"""Word (.docx) extractor — paragraphs + table cells, images ignored."""
import io

from docx import Document

from extractors.base import ExtractedContent, ExtractionError


def extract_docx(data: bytes, file_name: str) -> ExtractedContent:
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ExtractionError("This file doesn't look like a valid Word document, so it can't be read.") from e

    parts: list[str] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    # Pull table cell text too, so tabular content isn't lost.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise ExtractionError("No readable text was found in this Word document.")

    # Title: document core-properties title, else first line, else filename.
    title = ""
    try:
        if doc.core_properties.title:
            title = doc.core_properties.title.strip()
    except Exception:
        title = ""
    if not title:
        title = parts[0] if parts else file_name

    return ExtractedContent(title=title[:300], text=full_text)
